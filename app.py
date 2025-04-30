import os
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, jsonify
from werkzeug.utils import secure_filename
from deep_translator import GoogleTranslator
from docx import Document
import pandas as pd
import io
from reportlab.lib.pagesizes import letter
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from reportlab.lib.utils import ImageReader
# Inisialisasi aplikasi Flask
app = Flask(__name__)

# Tentukan folder upload dan ekstensi file yang diizinkan
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# output_dir = os.path.join(os.path.dirname(app.config['UPLOAD_FOLDER']), 'output')
# Direktori output
output_dir = os.path.join(app.root_path, 'static', 'output')

# Pastikan folder output ada
if not os.path.exists(output_dir):
    os.makedirs(output_dir)
           

# Fungsi untuk memeriksa ekstensi file yang diizinkan
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Fungsi untuk mengekstrak teks dari PDF
def extract_text_with_positions(input_pdf_path):
    doc = fitz.open(input_pdf_path)
    extracted = []
    page_sizes = []

    for page_num, page in enumerate(doc):
        rect = page.rect
        page_sizes.append((rect.width, rect.height))

        blocks = page.get_text("blocks")
        for block in blocks:
            x0, y0, x1, y1, text, block_no = block[:6]
            if text.strip():
                extracted.append({
                    "page": page_num,
                    "x": x0,
                    "y": y0,
                    "text": text.strip()
                })
    return extracted, page_sizes
def extract_lines(input_pdf_path):
    doc = fitz.open(input_pdf_path)
    lines = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        drawings = page.get_drawings()
        for d in drawings:
            if d['type'] == 'line':
                lines.append({
                    "page": page_num,
                    "p1": d['from'],
                    "p2": d['to']
                })

    return lines
def extract_images(input_pdf_path):
    doc = fitz.open(input_pdf_path)
    images = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images(full=True)
        for img in image_list:
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            bbox = page.get_image_bbox(img)
            images.append({
                "page": page_num,
                "bbox": bbox,  # (x0, y0, x1, y1)
                "image_bytes": image_bytes
            })

    return images
# Fungsi untuk menerjemahkan teks
def translate_doc(text, source_language='auto', target_language='id'):
    try:
        translated_text = GoogleTranslator(source=source_language, target=target_language).translate(text)
        return translated_text
    except Exception as e:
        return "[Translation failed]"

# Fungsi untuk menyimpan teks terjemahan ke dalam file Word
def save_to_word(text, output_word_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_word_path)

# Fungsi untuk menyimpan teks terjemahan ke dalam file Excel
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

def create_translated_pdf(text_blocks, page_sizes, output_pdf_path, lines, images,source_language,target_language):
    """Buat PDF baru dengan ukuran halaman sama seperti input, gambar garis, dan gambar"""

    c = canvas.Canvas(output_pdf_path)
    current_page = 0
    width, height = page_sizes[0]  # Ukuran halaman pertama
    c.setPageSize((width, height))
    c.setFont("Times-Roman", 10)  # Set font sebelum menggambar teks

    margin_left = 20
    margin_right = 20

    # Dapatkan garis-garis untuk halaman pertama
    page_lines = [line for line in lines if line['page'] == 0]
    draw_lines(c, page_lines, (width, height))

    # Dapatkan gambar untuk halaman pertama
    page_images = [img for img in images if img['page'] == 0]
    draw_images(c, page_images, (width, height))

    # Gambar teks dan terjemahan
    for block in text_blocks:
        if block['page'] != current_page:
            c.showPage()
            current_page = block['page']
            width, height = page_sizes[current_page]
            c.setPageSize((width, height))
            c.setFont("Times-Roman", 10)  # Wajib set font lagi setelah showPage()

            # Gambar garis dan gambar untuk halaman baru
            page_lines = [line for line in lines if line['page'] == current_page]
            draw_lines(c, page_lines, (width, height))
            page_images = [img for img in images if img['page'] == current_page]
            draw_images(c, page_images, (width, height))

        translated_text = translate_doc(block['text'],source_language,target_language)

        # Pastikan posisi teks memperhatikan margin
        x = max(block['x'], margin_left)
        y = height - block['y']

        textobject = c.beginText()
        textobject.setTextOrigin(x, y)
        textobject.setFont("Times-Roman", 10)

        # Hitung lebar maksimum untuk teks
        max_width = width - x - margin_right
        if max_width < 50:
            max_width = width - margin_left - margin_right  # kalau x terlalu dekat kanan, reset max_width

        # Bagi teks jadi beberapa baris supaya tidak melebar
        text_lines = simpleSplit(translated_text, "Times-Roman", 12, max_width)

        for text_line in text_lines:
            textobject.textLine(text_line)

        c.drawText(textobject)

    c.save()

def draw_lines(canvas_obj, lines, page_size):
    width, height = page_size

    for line in lines:
        x1, y1 = line['p1']
        x2, y2 = line['p2']
        canvas_obj.line(x1, height - y1, x2, height - y2)

def draw_images(canvas_obj, images, page_size):
    width, height = page_size

    for img in images:
        x0, y0, x1, y1 = img['bbox']
        img_width = x1 - x0
        img_height = y1 - y0

        y0_converted = height - y1  # Karena koordinat dibalik

        image = ImageReader(io.BytesIO(img['image_bytes']))
        canvas_obj.drawImage(image, x0, y0_converted, width=img_width, height=img_height)
# Fungsi untuk menerjemahkan teks menggunakan Google Translate
def translate_text_excel(text, target_lang='de'):
    try:
        hasil = GoogleTranslator(source='auto', target=target_lang).translate(text)
        return hasil
    except Exception as e:
        print(f"Terjadi kesalahan trans: {e}")
        return text 
def translate_column(df, target_lang='de'):
    for column in df.columns:
        df[column] = df[column].apply(lambda x: translate_text_excel(str(x), target_lang))
    return df
# Fungsi untuk mengekstrak teks dari file Word
def extract_text_from_word(docx_path):
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Fungsi untuk mengekstrak teks dari file Excel
def extract_text_from_excel(excel_path):
    df = pd.read_excel(excel_path)
    text = ""
    for index, row in df.iterrows():
        text += " ".join([str(cell) for cell in row]) + "\n"
    return text
def split_text(text, max_length=1000):
    return [text[i:i+max_length] for i in range(0, len(text), max_length)]
def process_translation(text, source_language='auto', target_language='id'):
    chunks = split_text(text)
    translated_chunks = [translate_doc(chunk, source_language, target_language) for chunk in chunks]
    translated_text = "".join(translated_chunks)
    return translated_text

# Route untuk halaman utama
@app.route('/')
def index():
    return render_template('index.html')
@app.route('/translateText', methods=['POST'])
def translate_text():
    data = request.get_json()

    source_lang = data['source_lang']
    target_lang = data['target_lang']
    text = data['text']

    try:
        # Menggunakan Deep Translate (GoogleTranslator sebagai contoh)
        translated_text = GoogleTranslator(source=source_lang, target=target_lang).translate(text)

        # Mengirimkan hasil terjemahan sebagai JSON response
        return jsonify({'translatedText': translated_text})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
# Route untuk menangani upload file
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'File tidak ditemukan'}), 400
    file = request.files['file']
    
    # Pastikan file memiliki nama dan ekstensi yang valid
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file_name_without_extension = os.path.splitext(filename)[0]
        file.save(file_path)
        print("Source Language:", request.form.get('source_language'))
        print("Target Language:", request.form.get('target_language'))
        # Mendapatkan bahasa sumber dan target dari form
        source_language = request.form.get('source_language', 'auto')  # 'auto' untuk deteksi otomatis
        target_language = request.form['target_language']
         # Periksa apakah target_language ada
        if not target_language:
            return jsonify({'error': 'Target language tidak disediakan.'}), 400
        if not source_language:
            return jsonify({'error': 'Target language tidak disediakan.'}), 400

        # Deteksi ekstensi file dan ekstrak teks
        
        if filename.lower().endswith(".pdf"):
        # Ekstraksi dan penerjemahan untuk PDF
            text_blocks, num_pages = extract_text_with_positions(file_path)
            # translated_text = process_translation(text, source_language, target_language)
            output_pdf_path = os.path.join(output_dir, f"terjemahan_{os.path.basename(filename)}")
            create_translated_pdf(text_blocks, num_pages, output_pdf_path,extract_lines(file_path),extract_images(file_path),source_language,target_language)
            print("Selesai! File disimpan di:", output_pdf_path)
            # save_to_pdf(translated_text, output_pdf_path)
            output_file_name = f"terjemahan_{file_name_without_extension}.pdf"
            link_text = "Unduh file PDF"

        elif filename.lower().endswith(".docx") :
            # Ekstraksi dan penerjemahan untuk Word
            text = extract_text_from_word(file_path)
            translated_text = process_translation(text, source_language, target_language)
            output_word_path = os.path.join(output_dir, f"terjemahan_{os.path.basename(filename)}")
            save_to_word(translated_text, output_word_path)
            output_file_name = f"terjemahan_{file_name_without_extension}.docx"
            link_text = "Unduh file Word"

        elif filename.lower().endswith(".xlsx"):
            # Ekstraksi dan penerjemahan untuk Excel
            df = pd.read_excel(file_path)# Menerjemahkan seluruh kolom
            df_translated = translate_column(df, target_language)
            output_excel_path = os.path.join(output_dir, f"terjemahan_{os.path.basename(filename)}")
            try:
                df_translated.to_excel(output_excel_path, index=False)
                print(f"File berhasil disimpan di: {output_excel_path}")
            except Exception as e:
                print(f"Terjadi kesalahan: {e}")


            output_file_name = f"terjemahan_{file_name_without_extension}.xlsx"
            print(output_file_name)
            print(output_excel_path)
            link_text = "Unduh file Excel"

        else:
            raise ValueError("File format not supported")
        

            # Mengembalikan response dengan link yang sesuai
        if output_file_name:
            return f"""File berhasil diupload dan diterjemahkan. 
                    <a href='static/output/{output_file_name}'>{link_text}</a>"""
        else:
            return "Terjadi kesalahan, file format tidak didukung."


    return jsonify({'error': 'File tidak valid, pastikan file berformat PDF, DOCX, atau XLSX.'}), 400


# Route untuk menangani pengunduhan file
@app.route('/static/output/<filename>')
def download_file(filename):
    print(filename)
    # """Mengunduh file yang ada di folder output"""
    return send_from_directory(output_dir, filename)

# Menjalankan aplikasi
if __name__ == '__main__':
    app.run(debug=True)
